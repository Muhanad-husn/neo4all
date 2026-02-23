"""
fixtures/generate_fixtures.py — Deterministic sample PDF and DOCX fixture generator.

Run once to produce fixtures/sample.pdf and fixtures/sample.docx, then commit
the resulting binary files.  The generator is kept alongside the fixtures so
the files can be regenerated if content changes are needed.

Determinism guarantees
----------------------
PDF:
  reportlab Canvas(invariant=1) suppresses live timestamps and uses a stable
  object-ID sequence so that identical content always produces identical bytes
  on the same reportlab version.

DOCX:
  python-docx document with core_properties timestamps pinned to a fixed date.
  The raw bytes are then passed through _normalise_zip() which rewrites every
  ZIP entry with a fixed timestamp (2024-01-01 00:00:00) and a lexicographically
  sorted entry order.  zlib compression is deterministic for fixed input +
  compression level, so the final file is byte-stable across repeated runs on
  the same python-docx / zlib version.

Runtime dependencies (dev only)
--------------------------------
  reportlab>=4.0        added to [project.optional-dependencies].dev
  python-docx>=1.1.0    added to [project.optional-dependencies].dev

Usage
-----
  python fixtures/generate_fixtures.py
"""

from __future__ import annotations

import io
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas

FIXTURES_DIR = Path(__file__).parent

# ---------------------------------------------------------------------------
# Shared text content — same corpus used in both fixtures for cross-format
# test coverage of the same logical document.
# ---------------------------------------------------------------------------

_INTRO_PARA = (
    "Knowledge graphs represent information as a network of entities and "
    "relationships, enabling machines to reason about complex domains. "
    "In a graph-based knowledge representation, nodes correspond to real-world "
    "objects such as organisations, people, locations, events, and concepts, "
    "while edges encode the semantic relationships that connect them. "
    "This structure supports a wide range of analytical tasks including entity "
    "resolution, link prediction, and question answering. "
    "Constructing a knowledge graph from unstructured text requires a pipeline "
    "that extracts named entities, resolves coreferences, and identifies "
    "relationships across sentences and documents. "
    "The quality of the resulting graph depends on the accuracy of each stage "
    "and on the schema that defines which entity and relationship types are "
    "considered relevant to the target domain."
)

_FIXTURE_NOTE = (
    "This fixture document is used to verify that the ingestion pipeline "
    "correctly parses document content, extracts text with page locators, and "
    "produces deterministic chunk identifiers from the page-locator hash function."
)

# Table content shared across both fixtures.
_TABLE_HEADERS_3COL = ("Entity Type", "Primary Property", "Example Values")
_TABLE_ROWS_3COL = [
    ("Organisation", "name",       "Acme Corp, OpenAI, WHO"),
    ("Person",       "full_name",  "Jane Smith, John Doe"),
    ("Location",     "place_name", "London, New York, Geneva"),
    ("Event",        "title",      "Annual Conference 2024"),
    ("Concept",      "label",      "Machine Learning, NLP"),
]

_TABLE_HEADERS_2COL = ("Entity Type", "Relationship Type")
_TABLE_ROWS_2COL = [
    ("Organisation", "EMPLOYS"),
    ("Person",       "LOCATED_IN"),
    ("Event",        "ORGANISED_BY"),
]


# ---------------------------------------------------------------------------
# PDF fixture
# ---------------------------------------------------------------------------


def _wrapped_text_lines(text: str, font: str, size: float, max_width: float,
                         c: canvas.Canvas) -> list[str]:
    """Split text into lines that fit within max_width for the given font/size."""
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip() if current else word
        if c.stringWidth(candidate, font, size) <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def _draw_paragraph(c: canvas.Canvas, x: float, y: float, width: float,
                    text: str, font: str = "Helvetica", size: float = 11.0,
                    leading: float = 16.0) -> float:
    """Draw word-wrapped paragraph text.  Returns y after the last line."""
    c.setFont(font, size)
    for line in _wrapped_text_lines(text, font, size, width, c):
        c.drawString(x, y, line)
        y -= leading
    return y


def _draw_table(c: canvas.Canvas, x: float, y: float,
                col_widths: list[float], row_height: float,
                headers: tuple[str, ...], rows: list[tuple[str, ...]]) -> float:
    """Draw a bordered table with a shaded header row.  Returns y below table."""
    header_bg = colors.HexColor("#4472C4")
    border_color = colors.HexColor("#CCCCCC")
    cell_pad_x = 0.2 * cm
    cell_pad_y = 0.2 * cm

    def _draw_row(row_data: tuple[str, ...], row_y: float, *, is_header: bool) -> float:
        cx = x
        for cell_text, col_w in zip(row_data, col_widths):
            if is_header:
                c.setFillColor(header_bg)
                c.rect(cx, row_y - row_height, col_w, row_height, fill=1, stroke=0)
                c.setFillColor(colors.white)
            else:
                c.setFillColor(colors.black)
            font = "Helvetica-Bold" if is_header else "Helvetica"
            c.setFont(font, 10)
            c.drawString(cx + cell_pad_x, row_y - row_height + cell_pad_y, cell_text)
            c.setStrokeColor(border_color)
            c.setLineWidth(0.5)
            c.rect(cx, row_y - row_height, col_w, row_height, fill=0, stroke=1)
            cx += col_w
        c.setFillColor(colors.black)
        return row_y - row_height

    current_y = _draw_row(headers, y, is_header=True)
    for row in rows:
        current_y = _draw_row(row, current_y, is_header=False)
    return current_y


def generate_pdf(dest: Path) -> None:
    """Generate fixtures/sample.pdf — deterministic 2-page PDF.

    Page 1: Title block, section heading, body paragraphs (~130 words).
    Page 2: "Sample Table" heading with a 3-column × 5-row data table.

    invariant=1 instructs reportlab to use a fixed creation timestamp and
    stable object-ID sequence, making the output byte-identical across
    repeated runs on the same reportlab version.
    """
    # invariant=1 → fixed creation date + stable object IDs (reportlab 4.x).
    c = canvas.Canvas(str(dest), pagesize=A4, invariant=1)
    width, height = A4
    margin = 2.5 * cm
    text_w = width - 2 * margin

    # Explicit metadata — belt-and-suspenders alongside invariant=1.
    c.setTitle("Sample Document")
    c.setAuthor("neo4all-fixtures")
    c.setSubject("SPEC-03 test fixture")
    c.setCreator("neo4all fixtures/generate_fixtures.py")

    # ── Page 1: Title + body ─────────────────────────────────────────────────
    y = height - margin

    # Title
    c.setFont("Helvetica-Bold", 24)
    c.drawString(margin, y - 24, "Sample Document")

    # Meta line
    c.setFont("Helvetica", 11)
    c.setFillColor(colors.grey)
    c.drawString(margin, y - 44, "Neo4all Test Fixture  |  Version 1.0")
    c.setFillColor(colors.black)

    # Horizontal rule
    c.setStrokeColor(colors.lightgrey)
    c.setLineWidth(0.5)
    c.line(margin, y - 56, width - margin, y - 56)
    c.setStrokeColor(colors.black)

    # Section heading
    c.setFont("Helvetica-Bold", 14)
    c.drawString(margin, y - 80, "Background")

    # Main body paragraph
    y2 = _draw_paragraph(c, margin, y - 100, text_w, _INTRO_PARA)

    # Second paragraph
    _draw_paragraph(c, margin, y2 - 10, text_w, _FIXTURE_NOTE)

    c.showPage()

    # ── Page 2: Table ────────────────────────────────────────────────────────
    y = height - margin

    # Page heading
    c.setFont("Helvetica-Bold", 16)
    c.drawString(margin, y - 16, "Sample Table")

    # Caption (matches spec: "caption 'Sample Table'")
    c.setFont("Helvetica-Oblique", 10)
    c.setFillColor(colors.grey)
    c.drawString(
        margin, y - 34,
        "Table 1 — Entity types extracted from knowledge graph literature",
    )
    c.setFillColor(colors.black)

    # 3-column × 5-row table
    col_widths = [5.0 * cm, 5.5 * cm, 6.5 * cm]
    table_bottom = _draw_table(
        c,
        x=margin,
        y=y - 55,
        col_widths=col_widths,
        row_height=0.8 * cm,
        headers=_TABLE_HEADERS_3COL,
        rows=[tuple(r) for r in _TABLE_ROWS_3COL],
    )

    # Footer note
    c.setFont("Helvetica", 9)
    c.setFillColor(colors.grey)
    c.drawString(margin, table_bottom - 16, "Source: neo4all fixtures — for testing only")
    c.setFillColor(colors.black)

    c.showPage()
    c.save()

    size = dest.stat().st_size
    print(f"  sample.pdf  -> {dest}  ({size:,} bytes)")


# ---------------------------------------------------------------------------
# DOCX fixture
# ---------------------------------------------------------------------------


def _normalise_zip(raw: bytes) -> bytes:
    """Rewrite a ZIP archive with fixed timestamps and sorted entry order.

    python-docx records the current system time on every ZIP entry.  This
    function re-packs all entries with a fixed date (2024-01-01 00:00:00)
    in lexicographic filename order so the output is byte-stable across
    repeated runs on the same python-docx / zlib version.

    Args:
        raw: Raw bytes of a DOCX (ZIP) file as returned by Document.save().

    Returns:
        DOCX bytes with normalised ZIP metadata.
    """
    FIXED_DATE = (2024, 1, 1, 0, 0, 0)
    src = zipfile.ZipFile(io.BytesIO(raw))
    out = io.BytesIO()

    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED, compresslevel=6) as zf:
        for info in sorted(src.infolist(), key=lambda i: i.filename):
            data = src.read(info.filename)
            new_info = zipfile.ZipInfo(info.filename, date_time=FIXED_DATE)
            new_info.compress_type = zipfile.ZIP_DEFLATED
            zf.writestr(new_info, data)

    return out.getvalue()


def generate_docx(dest: Path) -> None:
    """Generate fixtures/sample.docx — deterministic 2-section DOCX.

    Section 1: Heading "Introduction", two body paragraphs (~145 words total).
    Section 2: Heading "Data", intro sentence, 2-column × 4-row table
               (1 header row + 3 data rows).

    Core properties timestamps are pinned to 2024-01-01. The saved bytes are
    post-processed by _normalise_zip() to fix ZIP entry timestamps and sort
    order, making the file byte-stable across repeated runs.
    """
    doc = Document()

    # Pin all timestamps so they never vary between runs.
    _fixed = datetime(2024, 1, 1, 0, 0, 0)
    doc.core_properties.created = _fixed
    doc.core_properties.modified = _fixed
    doc.core_properties.author = "neo4all-fixtures"
    doc.core_properties.title = "Sample Document"
    doc.core_properties.subject = "SPEC-03 test fixture"

    # ── Section 1: Introduction ──────────────────────────────────────────────
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(_INTRO_PARA)
    doc.add_paragraph(_FIXTURE_NOTE)

    # ── Section 2: Data ──────────────────────────────────────────────────────
    doc.add_heading("Data", level=1)
    doc.add_paragraph(
        "The following table lists sample entity-relationship pairs used to "
        "exercise the schema validation and candidate generation logic."
    )

    # 2 columns × 4 rows (1 header + 3 data).
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    for col_idx, hdr_text in enumerate(_TABLE_HEADERS_2COL):
        run = hdr[col_idx].paragraphs[0].add_run(hdr_text)
        run.bold = True
        run.font.size = Pt(11)

    # Data rows
    for row_idx, (col1, col2) in enumerate(_TABLE_ROWS_2COL, start=1):
        cells = table.rows[row_idx].cells
        cells[0].text = col1
        cells[1].text = col2

    # Serialise → normalise ZIP metadata → write to disk.
    buf = io.BytesIO()
    doc.save(buf)
    stable_bytes = _normalise_zip(buf.getvalue())
    dest.write_bytes(stable_bytes)

    size = dest.stat().st_size
    print(f"  sample.docx -> {dest}  ({size:,} bytes)")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("Generating test fixtures…")
    generate_pdf(FIXTURES_DIR / "sample.pdf")
    generate_docx(FIXTURES_DIR / "sample.docx")
    print("Done.")
