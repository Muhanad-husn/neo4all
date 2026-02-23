"""
api/services/ingestion.py — Three-tier document parser fallback chain
(SPEC-03 S-03.1).

IngestorService parses uploaded documents using a cascading strategy:

  Tier 1 — Docling (primary):
    Full structural parsing: titles, paragraphs, tables, images, captions.
    Preferred for highest-fidelity extraction.  Output exported as Markdown.

  Tier 2 — Unstructured (secondary):
    Activated when Docling raises any exception or returns empty output.
    Broad format support; lower structural fidelity than Docling.

  Tier 3 — Raw text (tertiary):
    Activated when Unstructured also fails.  PyPDF2 for PDF files; python-docx
    for DOCX files.  Flat text only — no structural metadata.  Every chunk
    produced from a raw-fallback parse carries ChunkQualityFlag.raw_fallback.

  Hard reject:
    If all enabled tiers fail, IngestorService raises IngestionError.
    The system never silently accepts an empty or partial parse.

Manifest-based incremental reruns (SPEC-03 S-03.4 / SKILL-D R-D8):
    Before parsing, ingest_document() checks for an existing DocumentManifest
    via ArtifactsService.retrieve_manifest() (Redis → S3 fallback).  If a
    manifest is found whose content_hash and parser_config_hash both match the
    current document and parser configuration, the parse is skipped entirely.
    The caller (router) must check IngestResult.existing_manifest: when set,
    chunks already exist in Qdrant and S3; the pipeline step is a no-op.

    After chunking, callers invoke finalize_manifest() to persist the new
    manifest to S3 and cache it in Redis (TTL = 3600 s, SKILL-D R-D8).

Fallback cascade logging (SKILL-D R-D3):
    manifest_cache_hit   INFO    — emitted when a valid cached manifest allows
                                   skipping the parse step.
    manifest_cache_miss  INFO    — emitted when re-parsing is needed (reason
                                   documented in the structured fields).
    parser_selected      INFO    — emitted at the start of each tier attempt.
    parser_failed        WARNING — emitted when a tier raises or returns empty.
    fallback_triggered   WARNING — emitted on the transition between tiers.
    ingestion_complete   INFO    — emitted on success with doc_id, parser_used.
    ingestion_failed     ERROR   — emitted when all tiers are exhausted.

Parser toggles:
    ENABLE_DOCLING, ENABLE_UNSTRUCTURED, ENABLE_RAW_FALLBACK (api/config.py).
    Defaults to True; set False to disable a tier (e.g. in CI environments
    where a library is not installed).

Sensitive data:
    Raw document bytes and extracted text are NEVER logged (SKILL-D R-D5).
    Log doc_id and file_size_bytes instead.

Return type:
    IngestResult (Pydantic model, SKILL-A R-A3) carries the typed Document
    record, the extracted text for the chunking service, any default quality
    flags that must be applied to every chunk, and optionally an existing
    DocumentManifest (set when the parse step was skipped).
"""

from __future__ import annotations

import asyncio
import hashlib
import os
import tempfile
from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
from typing import TYPE_CHECKING, Any, Callable

from pydantic import BaseModel

from api.models.document import (
    ChunkQualityFlag,
    Document,
    DocumentManifest,
    derive_content_hash,
    derive_parser_config_hash,
)
from api.observability.logger import get_logger

if TYPE_CHECKING:
    from api.storage.artifacts import ArtifactsService

logger = get_logger(__name__)


# ---------------------------------------------------------------------------
# Public exception
# ---------------------------------------------------------------------------

class IngestionError(Exception):
    """Raised when all enabled parser tiers fail for a document.

    Callers (routers, ARQ jobs) catch this and emit a structured error
    response.  Never silently swallowed.

    Attributes:
        code:    Machine-readable error code for structured API responses
                 (e.g. "all_parsers_failed", "no_parsers_enabled").
        message: Human-readable description of the failure.
    """

    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code
        self.message = message


# ---------------------------------------------------------------------------
# Public return type (SKILL-A R-A3 — typed at module boundary)
# ---------------------------------------------------------------------------

class IngestResult(BaseModel):
    """Typed return value from IngestorService.ingest_document().

    Carries all information the calling layer needs to proceed:
      - document:             Immutable Document record (stored to S3/cache).
      - extracted_text:       Full text from the winning parser tier.  Passed
                              directly to ChunkingService.  NEVER logged.
                              Empty string when existing_manifest is set.
      - default_quality_flags: Applied to every Chunk from this document.
                              Non-empty only when the raw_text tier won.
      - existing_manifest:    Set when the parse was skipped because a valid
                              cached manifest was found (content_hash and
                              parser_config_hash both matched).  When set,
                              chunks already exist in Qdrant and S3; the caller
                              must skip chunking/indexing/storage steps.
                              None on fresh ingestion.

    This model exists here rather than in api/models/ because it is specific
    to the ingestion service output contract.  If it is consumed by 3+ modules
    it should be moved to api/models/ (SKILL-B R-B7).
    """

    document: Document
    extracted_text: str
    default_quality_flags: tuple[ChunkQualityFlag, ...] = ()
    existing_manifest: DocumentManifest | None = None


# ---------------------------------------------------------------------------
# Internal: deterministic doc_id derivation
# (Mirrors Document.doc_id computed field — used before constructing Document
#  so the cache key can be built from content_hash alone.)
# ---------------------------------------------------------------------------

def _derive_doc_id(run_id: str, source_identity: str, content_hash: str) -> str:
    """Compute doc_id without constructing a Document instance.

    Replicates Document.doc_id (CLAUDE.md §5):
        SHA-256(run_id + NUL + source_identity + NUL + content_hash)

    Null-byte separators prevent boundary-collision attacks (see document.py).
    Called once per ingest_document() invocation, before any cache or S3 I/O.
    """
    raw = f"{run_id}\x00{source_identity}\x00{content_hash}".encode("utf-8")
    return hashlib.sha256(raw).hexdigest()


# ---------------------------------------------------------------------------
# Internal: file type detection
# ---------------------------------------------------------------------------

def _detect_file_type(source_identity: str, file_bytes: bytes) -> str:
    """Return 'pdf', 'docx', or 'unknown'.

    Checks the filename extension first (fast path), then falls back to
    magic-byte detection for files whose extension is absent or non-standard.
    """
    suffix = Path(source_identity).suffix.lower()
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    # Magic bytes
    if file_bytes[:4] == b"%PDF":
        return "pdf"
    if file_bytes[:4] == b"PK\x03\x04":
        # ZIP-based Office format — best-effort guess is DOCX
        return "docx"
    return "unknown"


# ---------------------------------------------------------------------------
# Internal: parser tier implementations (synchronous — run via to_thread)
# ---------------------------------------------------------------------------

def _parse_with_docling(file_bytes: bytes, source_identity: str) -> str:
    """Tier 1: parse with Docling and export full-fidelity Markdown.

    Prefers an in-memory DocumentStream (Docling 2.x) to avoid any filesystem
    I/O.  If the DocumentStream import fails (older Docling version), falls
    back to a secure temp file that is deleted in the finally block.

    Args:
        file_bytes:      Raw document bytes.  Not logged (SKILL-D R-D5).
        source_identity: Used as the stream name for Docling metadata.

    Returns:
        Non-empty Markdown string.

    Raises:
        Exception: ImportError if Docling is not installed; any Docling
                   conversion error; ValueError if output is empty.
    """
    from docling.document_converter import DocumentConverter  # type: ignore[import-untyped]

    result: Any

    try:
        # Docling 2.x in-memory path — no filesystem I/O.
        from docling.datamodel.base_models import DocumentStream  # type: ignore[import-untyped]

        source: Any = DocumentStream(name=source_identity, stream=BytesIO(file_bytes))
        result = DocumentConverter().convert(source)

    except ImportError:
        # Older Docling — write to a temp file.  Use mkstemp for security
        # and close the fd via fdopen before Docling opens the path on Windows.
        suffix = Path(source_identity).suffix.lower() or ".tmp"
        tmp_fd, tmp_path = tempfile.mkstemp(suffix=suffix)
        try:
            with os.fdopen(tmp_fd, "wb") as fh:
                fh.write(file_bytes)
            # fd is now closed; Docling can open the path on Windows.
            result = DocumentConverter().convert(tmp_path)
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

    text: str = result.document.export_to_markdown()
    if not text.strip():
        raise ValueError("docling produced empty output")
    return text


def _parse_with_unstructured(file_bytes: bytes, source_identity: str) -> str:
    """Tier 2: parse with Unstructured via auto-partition.

    Accepts BytesIO directly; no filesystem I/O required.

    Args:
        file_bytes:      Raw document bytes.  Not logged (SKILL-D R-D5).
        source_identity: Passed as metadata_filename for element annotation.

    Returns:
        Non-empty text string with double-newline paragraph separation.

    Raises:
        Exception: ImportError if Unstructured is not installed; any
                   partition error; ValueError if output is empty.
    """
    from unstructured.partition.auto import partition  # type: ignore[import-untyped]

    elements = partition(
        file=BytesIO(file_bytes),
        metadata_filename=source_identity,
    )
    text = "\n\n".join(str(e) for e in elements if str(e).strip())
    if not text.strip():
        raise ValueError("unstructured produced empty output")
    return text


def _parse_with_raw_text(file_bytes: bytes, source_identity: str) -> str:
    """Tier 3: extract flat text with PyPDF2 (PDF) or python-docx (DOCX).

    This tier produces no structural metadata.  Every chunk derived from
    its output must carry ChunkQualityFlag.raw_fallback — that flag is set
    by IngestorService, not here.

    Args:
        file_bytes:      Raw document bytes.  Not logged (SKILL-D R-D5).
        source_identity: Used only for file-type detection and error messages.

    Returns:
        Non-empty flat text string.

    Raises:
        ValueError:  Unsupported file type, or extraction produced empty text.
        Exception:   ImportError if the required library is not installed;
                     any library-level parse error.
    """
    file_type = _detect_file_type(source_identity, file_bytes)

    if file_type == "pdf":
        import PyPDF2  # type: ignore[import-untyped]

        reader = PyPDF2.PdfReader(BytesIO(file_bytes))
        pages = [
            reader.pages[i].extract_text() or ""
            for i in range(len(reader.pages))
        ]
        text = "\n\n".join(p for p in pages if p.strip())

    elif file_type == "docx":
        import docx  # type: ignore[import-untyped]

        doc = docx.Document(BytesIO(file_bytes))
        text = "\n\n".join(
            para.text for para in doc.paragraphs if para.text.strip()
        )

    else:
        raise ValueError(
            f"unsupported file type for raw fallback: {source_identity!r}. "
            "Expected a .pdf or .docx file."
        )

    if not text.strip():
        raise ValueError(
            f"raw text extraction produced empty output for {file_type!r} file"
        )
    return text


# ---------------------------------------------------------------------------
# IngestorService
# ---------------------------------------------------------------------------

class IngestorService:
    """Parses documents through the three-tier fallback chain.

    Also coordinates manifest-based incremental reruns:
    - ingest_document() checks for an existing valid manifest before parsing.
    - finalize_manifest() persists a completed manifest to S3 + Redis cache.

    The service is stateless — parser functions carry no shared state.
    Multiple instances are safe to run concurrently.

    Args:
        settings:          Application settings (injectable for testing).
                           Defaults to the process-level singleton.
        artifacts_service: ArtifactsService for manifest retrieval and storage
                           (injectable for testing).  Defaults to the
                           process-level singleton via lazy import.
    """

    def __init__(
        self,
        settings: Any = None,
        artifacts_service: ArtifactsService | None = None,
    ) -> None:
        from api.config import get_settings

        self._settings = settings or get_settings()
        # Late-bound to avoid circular import at module load time.
        self._artifacts: ArtifactsService | None = artifacts_service

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _build_parser_config(self) -> dict[str, Any]:
        """Build the canonical parser config dict for parser_config_hash.

        Captures every setting that, if changed, would produce different
        chunk output from the same raw bytes.  The dict is serialised to a
        deterministic JSON string by derive_parser_config_hash() in
        api/models/document.py.

        Future: add parser library version strings when available at runtime.
        """
        return {
            "enable_docling": self._settings.ENABLE_DOCLING,
            "enable_unstructured": self._settings.ENABLE_UNSTRUCTURED,
            "enable_raw_fallback": self._settings.ENABLE_RAW_FALLBACK,
        }

    # ------------------------------------------------------------------
    # Public interface
    # ------------------------------------------------------------------

    async def ingest_document(
        self,
        run_id: str,
        file_bytes: bytes,
        source_identity: str,
    ) -> IngestResult:
        """Parse a document and return a typed IngestResult.

        Checks for a valid cached manifest before parsing (incremental reruns):
          - Derives content_hash and doc_id deterministically from the raw bytes.
          - Calls ArtifactsService.retrieve_manifest() (Redis → S3 fallback).
          - If a manifest is found whose content_hash AND parser_config_hash
            match the current document + configuration:
              · Logs "manifest_cache_hit" at INFO.
              · Returns immediately with IngestResult.existing_manifest set.
              · The caller must skip chunking, indexing, and storage steps.
          - If the manifest is stale (content or config changed):
              · Logs "manifest_cache_miss" at INFO with the change reason.
              · Falls through to the parser fallback chain.
          - If no manifest exists (first ingestion):
              · Logs "manifest_cache_miss" at DEBUG.
              · Falls through to the parser fallback chain.

        Args:
            run_id:          Governed run identifier — used in all log events
                             and participates in doc_id derivation.
            file_bytes:      Raw document bytes.  Never logged (SKILL-D R-D5).
            source_identity: Stable, content-independent source identifier
                             (filename, S3 key, canonical URL).

        Returns:
            IngestResult with the typed Document, extracted text, default
            quality flags, and (on cache hit) the existing DocumentManifest.

        Raises:
            IngestionError: If all enabled parser tiers fail, or if no tiers
                            are enabled in configuration.

        Log events (SKILL-D R-D3):
            manifest_cache_hit   INFO    — run_id, doc_id, chunk_count
            manifest_cache_miss  INFO    — run_id, doc_id, reason
                                           (content_changed | parser_config_changed)
            manifest_cache_miss  DEBUG   — run_id, doc_id, reason=not_found
            parser_selected      INFO    — parser, source_identity, tier_index, run_id
            parser_failed        WARNING — parser, reason, source_identity, run_id
            fallback_triggered   WARNING — from_parser, to_parser, reason, run_id
            ingestion_complete   INFO    — doc_id, parser_used, content_hash,
                                          file_size_bytes, run_id
            ingestion_failed     ERROR   — reason, last_parser, source_identity, run_id
        """
        content_hash = derive_content_hash(file_bytes)
        doc_id = _derive_doc_id(run_id, source_identity, content_hash)
        parser_config_hash = derive_parser_config_hash(self._build_parser_config())

        # --- Manifest cache / S3 check for incremental reruns ---
        if self._artifacts is not None:
            existing = await self._artifacts.retrieve_manifest(run_id, doc_id)
            if existing is not None:
                if (
                    existing.content_hash == content_hash
                    and existing.parser_config_hash == parser_config_hash
                ):
                    # Full hit: content unchanged and parser config matches.
                    # Skip parsing entirely — chunks already exist downstream.
                    logger.info(
                        "manifest_cache_hit",
                        run_id=run_id,
                        doc_id=doc_id,
                        chunk_count=len(existing.chunk_ids),
                    )
                    document = Document(
                        run_id=run_id,
                        source_identity=source_identity,
                        content_hash=content_hash,
                        # Sentinel: actual parser_used is from the original run.
                        # Callers should prefer existing_manifest.chunk_ids over
                        # document.parser_used when existing_manifest is set.
                        parser_used="cached",
                        created_at=existing.timestamp,
                    )
                    return IngestResult(
                        document=document,
                        extracted_text="",
                        existing_manifest=existing,
                    )

                # Stale manifest: determine which hash changed.
                if existing.content_hash != content_hash:
                    stale_reason = "content_changed"
                else:
                    stale_reason = "parser_config_changed"

                logger.info(
                    "manifest_cache_miss",
                    run_id=run_id,
                    doc_id=doc_id,
                    reason=stale_reason,
                )
            else:
                # No manifest found — first ingestion for this document.
                logger.debug(
                    "manifest_cache_miss",
                    run_id=run_id,
                    doc_id=doc_id,
                    reason="not_found",
                )

        # --- Parser fallback chain ---
        text, parser_used = await self._run_fallback_chain(
            run_id=run_id,
            file_bytes=file_bytes,
            source_identity=source_identity,
        )

        quality_flags: tuple[ChunkQualityFlag, ...] = (
            (ChunkQualityFlag.raw_fallback,) if parser_used == "raw_text" else ()
        )

        document = Document(
            run_id=run_id,
            source_identity=source_identity,
            content_hash=content_hash,
            parser_used=parser_used,
            created_at=datetime.now(UTC),
        )

        logger.info(
            "ingestion_complete",
            run_id=run_id,
            doc_id=document.doc_id,
            parser_used=parser_used,
            content_hash=content_hash,
            file_size_bytes=len(file_bytes),
        )

        return IngestResult(
            document=document,
            extracted_text=text,
            default_quality_flags=quality_flags,
        )

    async def finalize_manifest(self, manifest: DocumentManifest) -> str:
        """Persist a completed manifest to S3 and cache it in Redis.

        Called by the router after ChunkingService returns chunk_ids and the
        DocumentManifest is fully populated.  Delegates to
        ArtifactsService.store_manifest() which writes to S3 then caches with
        a 1-hour TTL (SKILL-D R-D8).

        Args:
            manifest: Completed DocumentManifest (chunk_ids fully populated).

        Returns:
            S3 key of the stored manifest (for audit trail).

        Raises:
            IngestionError: If ArtifactsService is not configured.
            StorageError:   Propagated from ArtifactsService.store_manifest()
                            on S3 failure.
        """
        if self._artifacts is None:
            raise IngestionError(
                code="no_artifacts_service",
                message=(
                    "ArtifactsService is not configured on IngestorService. "
                    "Inject an ArtifactsService instance at construction time."
                ),
            )
        return await self._artifacts.store_manifest(
            manifest.run_id, manifest.doc_id, manifest
        )

    # ------------------------------------------------------------------
    # Internal: fallback cascade
    # ------------------------------------------------------------------

    async def _run_fallback_chain(
        self,
        run_id: str,
        file_bytes: bytes,
        source_identity: str,
    ) -> tuple[str, str]:
        """Attempt each enabled parser tier in order and return on first success.

        Returns:
            (extracted_text, parser_used) — both non-empty strings.

        Raises:
            IngestionError: When all enabled tiers are exhausted, or when no
                            tiers are enabled in configuration.
        """
        s = self._settings

        # Build the ordered list of enabled (name, function) pairs.
        tiers: list[tuple[str, Callable[[bytes, str], str]]] = []
        if s.ENABLE_DOCLING:
            tiers.append(("docling", _parse_with_docling))
        if s.ENABLE_UNSTRUCTURED:
            tiers.append(("unstructured", _parse_with_unstructured))
        if s.ENABLE_RAW_FALLBACK:
            tiers.append(("raw_text", _parse_with_raw_text))

        if not tiers:
            logger.error(
                "ingestion_failed",
                run_id=run_id,
                source_identity=source_identity,
                reason="no_parsers_enabled",
            )
            raise IngestionError(
                code="no_parsers_enabled",
                message=(
                    "All parser tiers are disabled in configuration. "
                    "Enable at least one of ENABLE_DOCLING, ENABLE_UNSTRUCTURED, "
                    "or ENABLE_RAW_FALLBACK."
                ),
            )

        last_parser: str = ""
        last_reason: str = ""

        for idx, (parser_name, parse_fn) in enumerate(tiers):
            # Log the fallback transition before announcing the new tier.
            if idx > 0:
                logger.warning(
                    "fallback_triggered",
                    run_id=run_id,
                    from_parser=last_parser,
                    to_parser=parser_name,
                    reason=last_reason,
                    source_identity=source_identity,
                )

            logger.info(
                "parser_selected",
                run_id=run_id,
                parser=parser_name,
                source_identity=source_identity,
                tier_index=idx,
            )

            try:
                text = await asyncio.to_thread(parse_fn, file_bytes, source_identity)
                return text, parser_name
            except Exception as exc:
                last_reason = str(exc)
                last_parser = parser_name
                logger.warning(
                    "parser_failed",
                    run_id=run_id,
                    parser=parser_name,
                    source_identity=source_identity,
                    reason=last_reason,
                )

        # All tiers exhausted.
        logger.error(
            "ingestion_failed",
            run_id=run_id,
            source_identity=source_identity,
            reason="all_parsers_failed",
            last_parser=last_parser,
            last_reason=last_reason,
        )
        raise IngestionError(
            code="all_parsers_failed",
            message=(
                f"All {len(tiers)} enabled parser tier(s) failed for "
                f"{source_identity!r}. "
                f"Last error from {last_parser!r}: {last_reason}"
            ),
        )


# ---------------------------------------------------------------------------
# Process-level singleton factory
# ---------------------------------------------------------------------------

_service_instance: IngestorService | None = None


def get_ingestor_service() -> IngestorService:
    """Return the process-level IngestorService singleton.

    Uses a module-level variable so that tests can replace the instance
    without import-level side effects (same pattern as get_schema_service).

    ArtifactsService is wired in here to keep the import deferred (avoids
    any circular-import risk at module load time).
    """
    global _service_instance
    if _service_instance is None:
        from api.storage.artifacts import get_artifacts_service

        _service_instance = IngestorService(
            artifacts_service=get_artifacts_service(),
        )
    return _service_instance
